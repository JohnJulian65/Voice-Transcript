import sounddevice as sd
import numpy as np
import scipy.io.wavfile as wav
import requests
import datetime
import os
import ssl
import logging
import re
from urllib3.util.retry import Retry
from requests.adapters import HTTPAdapter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

# Set up logging
logging.basicConfig(
    filename='transcription.log',
    level=logging.INFO,
    format='%(asctime)s - %(levelname)s - %(message)s'
)

# Configuration
GROQ_API_KEY = "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"
API_ENDPOINT = "https://api.groq.com/openai/v1/audio/transcriptions"
MODEL = "whisper-large-v3-turbo"
LOG_FILE = "conversation_log.txt"
STRUCTURED_LOG_FILE = "structured_conversation_log.txt"
TEMP_AUDIO_FILE = "temp_audio.wav"
DURATION = 30  # seconds
SAMPLE_RATE = 16000  # recommended by Groq for optimal performance
SPEAKER_MEMORY_FILE = "speaker_memory.txt"  # To store identified speakers

HEADERS = {
    "Authorization": f"Bearer {GROQ_API_KEY}"
}

# Known roles and titles that may appear in the hearing
ROLES = [
    "chairman", "chairwoman", "chair", "ranking member", "representative",
    "senator", "ambassador", "mr.", "ms.", "mrs.", "dr.", "secretary", 
    "commissioner", "director", "administrator", "congressman", "congresswoman",
    "subcommittee", "committee"
]

# Track current and previous speakers
current_speaker = None
previous_speaker = None
speaker_memory = {}  # Dictionary to store speaker references
last_speaker_change_time = datetime.datetime.now()

def initialize_speaker_memory():
    """Initialize or load the speaker memory from file"""
    global speaker_memory
    if os.path.exists(SPEAKER_MEMORY_FILE):
        with open(SPEAKER_MEMORY_FILE, "r", encoding="utf-8") as f:
            lines = f.readlines()
            for line in lines:
                if ":" in line:
                    key, value = line.strip().split(":", 1)
                    speaker_memory[key.strip()] = value.strip()
    
    # Initialize with common references
    default_references = {
        "chairman": "Committee Chairman",
        "ranking member": "Ranking Member",
        "senator from rhode island": "Senator from Rhode Island",
        "senator cramer": "Senator Cramer"
    }
    
    for key, value in default_references.items():
        if key not in speaker_memory:
            speaker_memory[key] = value

def save_speaker_memory():
    """Save the speaker memory to file"""
    with open(SPEAKER_MEMORY_FILE, "w", encoding="utf-8") as f:
        for key, value in speaker_memory.items():
            f.write(f"{key}: {value}\n")

def record_audio(duration, sample_rate):
    print(f"Recording audio for {duration} seconds...")
    audio = sd.rec(int(duration * sample_rate), samplerate=sample_rate, channels=1, dtype='int16', device=1)
    sd.wait()
    return audio

def save_audio(audio_data, sample_rate, filename):
    wav.write(filename, sample_rate, audio_data)
    print(f"Audio segment saved as {filename}")

def create_robust_session():
    """Create a requests session with retries for older requests versions"""
    session = requests.Session()
    
    retry_strategy = Retry(
        total=3,
        backoff_factor=1,
        status_forcelist=[500, 502, 503, 504, 429],
        allowed_methods=["POST"]
    )
    
    adapter = HTTPAdapter(max_retries=retry_strategy)
    session.mount("https://", adapter)
    
    # For corporate environments, you might need this:
    session.verify = True  # Keep SSL verification on
    return session

def transcribe_audio(filename):
    try:
        with open(filename, "rb") as audio_file:
            files = {"file": (filename, audio_file, "audio/wav")}
            data = {
                "model": MODEL,
                "response_format": "json",
                "language": "en"
            }
            session = create_robust_session()
            response = session.post(
                API_ENDPOINT,
                headers=HEADERS,
                files=files,
                data=data,
                timeout=30
            )
            response.raise_for_status()
            json_data = response.json()
            print("API Response:", json_data)
            return json_data.get('text', '')
    except requests.exceptions.SSLError as ssl_err:
        print(f"❌ SSL error occurred: {ssl_err}")
        return "[SSL Error: Transcription failed]"
    except requests.exceptions.RequestException as req_err:
        print(f"❌ Request error occurred: {req_err}")
        return "[Request Error: Transcription failed]"
    except Exception as e:
        print(f"❌ Unexpected error: {e}")
        return "[Unknown Error: Transcription failed]"

def extract_name_reference(text):
    """Extract references to people in the text"""
    # Patterns for speaker references
    reference_patterns = [
        # Direct introduction
        r'(?:recognized|recognizes|welcomes|introduces|yield to)\s+(?:the\s+)?(?:distinguished\s+)?((?:gentleman|gentlewoman|senator|representative|ambassador)\s+(?:from\s+\w+(?:\s+\w+)?|(?:\w+\s+)+))',
        # The senator/representative from [state]
        r'(?:the\s+)?((?:senator|representative)\s+from\s+(?:\w+\s*)+)',
        # Specific titles with names
        r'(?:(?:sub)?committee\s+chair(?:man|woman|person)?\s+(?:\w+\s+)+)',
        r'(?:ranking\s+member\s+(?:\w+\s+)+)',
        # Names with titles
        r'(?:senator|representative|ambassador|secretary|commissioner)\s+(\w+)'
    ]
    
    for pattern in reference_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            for match in matches:
                if isinstance(match, tuple):  # In case there are capturing groups
                    for m in match:
                        if m and len(m) > 3:  # Avoid short matches
                            return m.strip()
                elif match and len(match) > 3:  # Avoid short matches
                    return match.strip()
    
    return None

def update_speaker_memory(text):
    """Update the speaker memory based on text analysis"""
    global speaker_memory, current_speaker, previous_speaker, last_speaker_change_time
    
    # Look for introductions or references
    new_reference = extract_name_reference(text)
    if new_reference:
        # Store this reference for future use
        key = new_reference.lower()
        if key not in speaker_memory:
            speaker_memory[key] = new_reference
            save_speaker_memory()
            print(f"Added new speaker reference: {new_reference}")
        
        # Check if this is a new speaker
        if not current_speaker or new_reference.lower() != current_speaker.lower():
            previous_speaker = current_speaker
            current_speaker = new_reference
            last_speaker_change_time = datetime.datetime.now()
            return True, new_reference
    
    return False, None

def identify_speaker(text):
    """Identify the speaker from the transcript text"""
    global current_speaker, previous_speaker, last_speaker_change_time
    
    # First check for explicit speaker patterns
    explicit_patterns = [
        # Direct identification patterns: "Chairman Smith:" or "Mr. Johnson says"
        r'(?:^|\s)((?:' + '|'.join(ROLES) + r')\s+\w+(?:\s+\w+)*)(?:\s*:|says|\sstated)',
        # Simple name mentions that might be speakers
        r'(?:^|\s)(\w+\s+\w+)(?:\s*:)'
    ]
    
    for pattern in explicit_patterns:
        matches = re.findall(pattern, text, re.IGNORECASE)
        if matches:
            # Found an explicit speaker
            speaker = matches[0].strip()
            # Update current speaker if different
            if not current_speaker or speaker.lower() != current_speaker.lower():
                previous_speaker = current_speaker
                current_speaker = speaker
                last_speaker_change_time = datetime.datetime.now()
            return format_speaker_name(speaker)
    
    # Look for speaker references/introductions
    is_introduction, new_speaker = update_speaker_memory(text)
    if is_introduction and new_speaker:
        return format_speaker_name(new_speaker)
    
    # Check for question patterns that might indicate a switch to previous speaker
    question_indicators = extract_question_answer(text) == "question"
    if question_indicators and previous_speaker:
        # If we have a question and a previous speaker, assume they're asking
        temp = current_speaker
        current_speaker = previous_speaker
        previous_speaker = temp
        last_speaker_change_time = datetime.datetime.now()
        return format_speaker_name(current_speaker)
    
    # If we've had the same speaker for more than 2 segments, consider switching
    time_since_change = (datetime.datetime.now() - last_speaker_change_time).total_seconds()
    if time_since_change > DURATION * 2 and previous_speaker:
        # Assume speaker changed back to previous speaker
        temp = current_speaker
        current_speaker = previous_speaker
        previous_speaker = temp
        last_speaker_change_time = datetime.datetime.now()
        return format_speaker_name(current_speaker)
    
    # If we have a current speaker, continue with them
    if current_speaker:
        return format_speaker_name(current_speaker)
    
    # Default fallback
    return "Speaker"

def format_speaker_name(speaker):
    """Format the speaker name appropriately"""
    if not speaker:
        return "Speaker"
        
    # Capitalize role words
    formatted = speaker
    for role in ROLES:
        if role in formatted.lower():
            pattern = re.compile(re.escape(role), re.IGNORECASE)
            formatted = pattern.sub(role.capitalize(), formatted)
    
    # Capitalize proper names
    words = formatted.split()
    if len(words) > 1:
        formatted = " ".join([words[0]] + [w.capitalize() for w in words[1:]])
    
    # Check if it matches any stored reference
    lower_formatted = formatted.lower()
    for key, value in speaker_memory.items():
        if key in lower_formatted or lower_formatted in key:
            return value
            
    return formatted

def extract_question_answer(text):
    """Attempt to identify if text is a question or an answer"""
    # Question indicators
    question_patterns = [
        r'\?',
        r'^(?:what|how|why|where|when|is|are|do|does|can|could|would|should|will)',
        r'(?:can you|would you|could you|do you|have you|is there|are there)'
    ]
    
    for pattern in question_patterns:
        if re.search(pattern, text.lower()):
            return "question"
    
    return "answer"

def structure_transcript(text):
    """Structure the transcript into speaker sections"""
    # Try to identify speaker
    speaker = identify_speaker(text)
    
    # Format the transcript
    formatted_text = f"**{speaker}:** {text}"
    
    return formatted_text

def log_transcript(text, logfile):
    """Log the raw transcript to a file"""
    timestamp = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    with open(logfile, "a", encoding="utf-8") as log:
        log.write(f"[{timestamp}] {text}\n")
    print(f"Transcript logged at {timestamp}")

def log_structured_transcript(text, logfile):
    """Log the structured transcript to a file"""
    with open(logfile, "a", encoding="utf-8") as log:
        log.write(f"{text}\n\n")
    print("Structured transcript logged")

def create_structured_docx(input_txt_path, output_docx_path):
    """Create a structured Word document from the text file"""
    doc = Document()
    
    # Set the document style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)
    
    # Read the structured transcript
    with open(input_txt_path, "r", encoding="utf-8") as text_file:
        content = text_file.read()
    
    # Split by double newlines to separate speaker sections
    sections = content.split("\n\n")
    
    for section in sections:
        if not section.strip():
            continue
            
        paragraph = doc.add_paragraph()
        
        # Check if this is a speaker header
        speaker_match = re.search(r'^\*\*(.+?):\*\*', section)
        if speaker_match:
            # Split into speaker and text
            speaker_text = re.split(r'^\*\*(.+?):\*\*\s*', section, 1)
            if len(speaker_text) > 2:
                speaker = speaker_text[1]
                text = speaker_text[2]
                
                # Add speaker with bold formatting
                speaker_run = paragraph.add_run(f"{speaker}: ")
                speaker_run.bold = True
                
                # Add the rest of the text
                paragraph.add_run(text)
        else:
            # If no speaker format is found, just add the text
            paragraph.add_run(section)
    
    doc.save(output_docx_path)
    print(f"✅ Structured transcript saved to {output_docx_path}")

def cleanup(filename):
    """Remove temporary files"""
    if os.path.exists(filename):
        os.remove(filename)
        print(f"Cleaned up {filename}")

def main():
    print("🚀 Starting continuous transcription with improved speaker recognition. Press Ctrl+C to stop.")
    
    # Initialize speaker memory
    initialize_speaker_memory()
    
    # Initialize or clear the structured log file
    with open(STRUCTURED_LOG_FILE, "w", encoding="utf-8") as f:
        f.write("")
    
    try:
        while True:
            audio_data = record_audio(DURATION, SAMPLE_RATE)
            save_audio(audio_data, SAMPLE_RATE, TEMP_AUDIO_FILE)

            # Get raw transcript
            transcript = transcribe_audio(TEMP_AUDIO_FILE)
            
            if not transcript or transcript.startswith("["):
                print("Skipping empty or error transcript")
                continue
            
            # Log raw transcript (for backup)
            log_transcript(transcript, LOG_FILE)
            
            # Structure the transcript
            structured_transcript = structure_transcript(transcript)
            
            # Log structured transcript
            log_structured_transcript(structured_transcript, STRUCTURED_LOG_FILE)
            
            # Create structured Word document
            create_structured_docx(STRUCTURED_LOG_FILE, "structured_output.docx")
            
            # Optional: Clean up temporary audio file
            #cleanup(TEMP_AUDIO_FILE)

    except KeyboardInterrupt:
        print("\n❌ Transcription stopped by user.")
        # Save speaker memory before exiting
        save_speaker_memory()

if __name__ == "__main__":
    main()
